"""
Sushruta — Document Service
==============================

Business logic for medical document upload, retrieval, and management.

File handling:
- Files are saved to disk with UUID-based names (prevents collisions).
- Original filename is preserved in the database for display.
- File size and type are validated before saving.
- Text extraction runs synchronously in Phase 1 (async in Phase 2).

Text extraction:
- PDF: pypdf extracts text from all pages.
- DOCX: python-docx extracts paragraph text.
- Images: Deferred to Phase 2 (OCR with Tesseract or cloud API).

Processing pipeline:
  uploaded → processing → ready | failed

Security:
- Only the owning doctor can access documents.
- Patient ownership is verified before any document operation.
"""

import json
import os
import uuid

import aiofiles
from fastapi import HTTPException, UploadFile, status
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.core.audit import AuditAction, create_audit_log
from app.db.models import Doctor, Document
from app.schemas.document import (
    DocumentDetailResponse,
    DocumentListResponse,
    DocumentResponse,
)
from app.services.patient_service import get_patient_model

settings = get_settings()


def _extract_text_from_pdf(file_path: str) -> str | None:
    """
    Extract text from a PDF file using pypdf.

    Returns concatenated text from all pages, or None if extraction fails.
    pypdf handles most standard PDFs but may fail on scanned/image PDFs
    (those need OCR, added in Phase 2).
    """
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return "\n".join(text_parts) if text_parts else None
    except Exception:
        return None


def _extract_text_from_docx(file_path: str) -> str | None:
    """
    Extract text from a DOCX file using python-docx.

    Returns concatenated paragraph text, or None if extraction fails.
    """
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n".join(text_parts) if text_parts else None
    except Exception:
        return None


def _extract_text(file_path: str, file_type: str) -> str | None:
    """
    Route text extraction to the appropriate handler based on file type.

    Images (.png, .jpg, .jpeg) return None in Phase 1 — OCR is Phase 2.
    """
    if file_type == ".pdf":
        return _extract_text_from_pdf(file_path)
    elif file_type == ".docx":
        return _extract_text_from_docx(file_path)
    else:
        # Images: text extraction deferred to Phase 2 (OCR)
        return None


async def upload_document(
    db: AsyncSession,
    patient_id: int,
    file: UploadFile,
    doctor: Doctor,
    ip_address: str | None = None,
) -> DocumentResponse:
    """
    Upload a medical document for a patient.

    Flow:
    1. Verify patient ownership.
    2. Validate file type (extension check).
    3. Read file content and validate size.
    4. Generate UUID filename and save to disk.
    5. Extract text (PDF/DOCX).
    6. Create document record with processing status.
    7. Audit log.
    8. Commit and return.

    Raises
    ------
    HTTPException 415
        If file type is not in ALLOWED_FILE_TYPES.
    HTTPException 413
        If file exceeds MAX_FILE_SIZE_MB.
    HTTPException 404
        If patient not found or not owned by doctor.
    """
    # Step 1: Verify patient ownership
    await get_patient_model(db, patient_id, doctor)

    # Step 2: Validate file type
    original_filename = file.filename or "unknown"
    file_ext = os.path.splitext(original_filename)[1].lower()

    if file_ext not in settings.ALLOWED_FILE_TYPES:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"File type '{file_ext}' is not supported. "
                   f"Allowed: {', '.join(settings.ALLOWED_FILE_TYPES)}",
        )

    # Step 3: Read content and validate size
    content = await file.read()
    file_size = len(content)

    if file_size > settings.max_file_size_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_CONTENT_TOO_LARGE,
            detail=f"File size ({file_size / 1024 / 1024:.1f}MB) exceeds "
                   f"maximum allowed size ({settings.MAX_FILE_SIZE_MB}MB)",
        )

    # Step 4: Generate unique filename and save to disk
    stored_filename = f"{uuid.uuid4().hex}{file_ext}"
    upload_dir = os.path.join(settings.UPLOAD_DIR, str(patient_id))
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, stored_filename)

    async with aiofiles.open(file_path, "wb") as f:
        await f.write(content)

    # Step 5: Extract text
    extracted_text = _extract_text(file_path, file_ext)
    processing_status = "ready" if extracted_text else "uploaded"

    # For image files, status stays "uploaded" (OCR in Phase 2)
    # For PDFs/DOCX with no extractable text, mark as "failed"
    if file_ext in (".pdf", ".docx") and extracted_text is None:
        processing_status = "failed"

    # Step 6: Create document record
    document = Document(
        patient_id=patient_id,
        doctor_id=doctor.id,
        original_filename=original_filename,
        stored_filename=stored_filename,
        file_type=file_ext,
        file_size_bytes=file_size,
        extracted_text=extracted_text,
        processing_status=processing_status,
    )
    db.add(document)
    await db.flush()

    # Step 7: Audit log
    await create_audit_log(
        db,
        doctor_id=doctor.id,
        patient_id=patient_id,
        action=AuditAction.DOCUMENT_UPLOADED,
        resource_type="document",
        resource_id=document.id,
        details=json.dumps({
            "filename": original_filename,
            "file_type": file_ext,
            "size_bytes": file_size,
            "processing_status": processing_status,
        }),
        ip_address=ip_address,
    )

    # Step 8: Commit
    await db.commit()
    await db.refresh(document)
    return DocumentResponse.model_validate(document)


async def list_documents(
    db: AsyncSession,
    patient_id: int,
    doctor: Doctor,
) -> DocumentListResponse:
    """
    List all active documents for a patient.

    Verifies patient ownership before listing.
    Returns lightweight DocumentResponse (no extracted_text).
    """
    # Verify patient ownership
    await get_patient_model(db, patient_id, doctor)

    result = await db.execute(
        select(Document).where(
            Document.patient_id == patient_id,
            Document.doctor_id == doctor.id,
            Document.is_active == True,  # noqa: E712
        ).order_by(Document.created_at.desc())
    )
    documents = result.scalars().all()

    # Get total count
    count_result = await db.execute(
        select(func.count()).select_from(
            select(Document).where(
                Document.patient_id == patient_id,
                Document.doctor_id == doctor.id,
                Document.is_active == True,  # noqa: E712
            ).subquery()
        )
    )
    total = count_result.scalar() or 0

    return DocumentListResponse(
        documents=[DocumentResponse.model_validate(d) for d in documents],
        total=total,
    )


async def get_document(
    db: AsyncSession,
    patient_id: int,
    document_id: int,
    doctor: Doctor,
) -> DocumentDetailResponse:
    """
    Get a single document with full details (including extracted text).

    Verifies both patient and document ownership.

    Raises
    ------
    HTTPException 404
        If document not found, inactive, or belongs to another doctor.
    """
    # Verify patient ownership
    await get_patient_model(db, patient_id, doctor)

    result = await db.execute(
        select(Document).where(
            Document.id == document_id,
            Document.patient_id == patient_id,
            Document.doctor_id == doctor.id,
            Document.is_active == True,  # noqa: E712
        )
    )
    document = result.scalar_one_or_none()

    if document is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )

    return DocumentDetailResponse.model_validate(document)


async def delete_document(
    db: AsyncSession,
    patient_id: int,
    document_id: int,
    doctor: Doctor,
    ip_address: str | None = None,
) -> dict:
    """
    Soft-delete a document by setting is_active=False.

    The file remains on disk and the DB record is retained.
    In a production system, a background job would handle
    physical file cleanup based on retention policies.
    """
    # Verify patient ownership
    await get_patient_model(db, patient_id, doctor)

    result = await db.execute(
        select(Document).where(
            Document.id == document_id,
            Document.patient_id == patient_id,
            Document.doctor_id == doctor.id,
            Document.is_active == True,  # noqa: E712
        )
    )
    document = result.scalar_one_or_none()

    if document is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )

    document.is_active = False
    await db.flush()

    await create_audit_log(
        db,
        doctor_id=doctor.id,
        patient_id=patient_id,
        action=AuditAction.DOCUMENT_DELETED,
        resource_type="document",
        resource_id=document.id,
        details=json.dumps({"filename": document.original_filename}),
        ip_address=ip_address,
    )

    await db.commit()
    return {"message": "Document removed"}
